from __future__ import annotations

import html
import io
import json
import re
from collections.abc import Iterable, Sequence
from typing import Any
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models import SiteAnalysis
from app.sef.report import HumanReviewedReportV1


WORD_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
MARKDOWN_MEDIA_TYPE = "text/markdown"
REPORT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
TABLE_CELL_MARGIN_DXA = {
    "top": 80,
    "bottom": 80,
    "start": 120,
    "end": 120,
}
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x5B, 0x65, 0x73)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "EAF0F7"


def export_filename(report_id: str, extension: str) -> str:
    safe_id = re.sub(r"[^a-zA-Z0-9_-]", "_", report_id).strip("_")
    safe_id = safe_id or "report"
    return f"aimeton-{safe_id}.{extension}"


def _safe_http_url(value: Any) -> str | None:
    url = str(value).strip()
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return None
    return url


def _markdown_url(value: Any) -> str:
    url = _safe_http_url(value)
    return f"<{url}>" if url else _md_text(value)


def _display_value(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    if value is None:
        return "—"
    return str(value)


def _md_text(value: Any) -> str:
    text = html.escape(_display_value(value), quote=False)
    for marker in ("\\", "`", "*", "_", "{", "}", "[", "]", "<", ">"):
        text = text.replace(marker, f"\\{marker}")
    return text


def _md_cell(value: Any) -> str:
    return _md_text(value).replace("|", r"\|").replace("\r", "").replace(
        "\n",
        "<br>",
    )


def _md_bullets(values: Iterable[Any], empty: str = "Нет данных.") -> list[str]:
    rendered = [f"- {_md_text(value)}" for value in values]
    return rendered or [empty]


def _append_markdown_table(
    lines: list[str],
    headers: Sequence[str],
    rows: Iterable[Sequence[Any]],
) -> None:
    prepared = list(rows)
    lines.append("| " + " | ".join(_md_cell(item) for item in headers) + " |")
    lines.append("| " + " | ".join("---" for _ in headers) + " |")
    if not prepared:
        empty = ["Нет проверенных данных."] + ["—"] * (len(headers) - 1)
        prepared.append(empty)
    lines.extend(
        "| " + " | ".join(_md_cell(item) for item in row) + " |"
        for row in prepared
    )
    lines.append("")


def render_site_analysis_markdown(analysis: SiteAnalysis) -> str:
    opportunity = analysis.commercial_opportunity
    action = analysis.action_package
    lines = [
        "# AIMETON — предварительный анализ компании",
        "",
        "> Статус: предварительный результат старого контура анализа. "
        "Он не является подписанным Report v1 и не прошёл human sign-off.",
        "",
        f"**Компания:** {_md_text(analysis.company_name)}  ",
        f"**Сайт:** {_markdown_url(analysis.url)}",
        "",
        "## Состояние допустимости",
        "",
        f"- Analysis state: {analysis.readiness.analysis_state}",
        f"- Client release: "
        f"{str(analysis.readiness.client_release_eligible).lower()}",
        f"- УДП: {analysis.readiness.sufficiency_level}",
        f"- Identity: {analysis.readiness.identity_state}",
        f"- Полнота профиля: "
        f"{analysis.readiness.profile_completeness:.0%}",
        f"- Качество evidence: "
        f"{analysis.readiness.evidence_quality:.0%}",
        f"- Коммерческий приоритет: "
        f"{analysis.readiness.commercial_priority}/100",
        f"- Budget: {analysis.readiness.budget_state}",
        f"- Providers: "
        f"{', '.join(f'{key}={value}' for key, value in sorted(analysis.readiness.provider_states.items())) or 'not_reported'}",
        f"- Обязательные вертикали: "
        f"{', '.join(f'{item.code}={item.state}' for item in analysis.readiness.required_verticals) or 'not_reported'}",
        f"- Блокеры: "
        f"{', '.join(analysis.readiness.release_blockers) or '—'}",
        "",
        _md_text(analysis.business_summary),
        "",
        "## Коммерческая возможность",
        "",
        f"**Тип:** {_md_text(opportunity.opportunity_type)}  ",
        f"**Оценка:** {opportunity.score}/100  ",
        f"**Квалификация:** {_md_text(opportunity.qualification)}",
        "",
        f"**Гипотеза проблемы:** {_md_text(opportunity.problem_hypothesis)}",
        "",
        f"**Рекомендуемое решение:** "
        f"{_md_text(opportunity.recommended_solution)}",
        "",
        f"**Ожидаемая ценность:** {_md_text(opportunity.expected_value)}",
        "",
        "## Экономические сигналы",
        "",
    ]
    _append_markdown_table(
        lines,
        ("Сигнал", "Основание", "Эффект", "Уверенность", "Источники"),
        (
            (
                item.signal,
                item.evidence,
                item.business_effect,
                item.confidence,
                ", ".join(item.source_ids) or "—",
            )
            for item in analysis.economic_signals
        ),
    )
    lines.extend(["## Факты о компании", ""])
    _append_markdown_table(
        lines,
        ("Поле", "Значение", "Период", "Уверенность", "Источники", "Примечание"),
        (
            (
                item.field,
                item.value,
                item.period or "—",
                item.confidence,
                ", ".join(item.source_ids) or "—",
                item.note or "—",
            )
            for item in analysis.company_facts
        ),
    )
    lines.extend(["## Бизнес-машина AIMETON 4×4", ""])
    _append_markdown_table(
        lines,
        (
            "Код",
            "Оператор",
            "Вершина",
            "Результат",
            "Статус",
            "Источники",
            "Значение для продажи",
        ),
        (
            (
                item.code,
                item.detail_operator,
                item.vertex,
                item.finding,
                f"{item.status}; {item.confidence}",
                ", ".join(item.source_ids) or "—",
                item.sales_relevance or "—",
            )
            for item in analysis.business_machine_4x4
        ),
    )
    lines.extend(["## Подходящие AI-инструменты", ""])
    _append_markdown_table(
        lines,
        ("Инструмент", "Назначение", "Польза", "Приоритет"),
        (
            (item.name, item.purpose, item.benefit, item.priority)
            for item in analysis.agents
        ),
    )
    lines.extend(
        [
            "## Пакет действия",
            "",
            f"**Предполагаемый ЛПР:** "
            f"{_md_text(action.decision_maker_hypothesis)}",
            "",
            f"**Основание для контакта:** {_md_text(action.contact_reason)}",
            "",
            "### Демонстрационный сценарий",
            "",
            *_md_bullets(action.demo_scenario),
            "",
            "### Первое сообщение",
            "",
            *[
                f"> {_md_text(line)}"
                for line in action.first_message.splitlines() or [""]
            ],
            "",
            f"**Следующий шаг:** {_md_text(action.next_action)}",
            "",
            "## Источники",
            "",
        ]
    )
    for source in analysis.sources:
        lines.extend(
            [
                f"### {_md_text(source.id)} — {_md_text(source.title)}",
                "",
                f"- URL: {_markdown_url(source.url)}",
                f"- Проверено: {_md_text(source.accessed_at)}",
                f"- Тип: {_md_text(source.source_type)}",
                f"- Уровень: {_md_text(source.evidence_level)}",
                f"- Цитата: {_md_text(source.evidence_quote)}",
            ]
        )
        if source.document_digest:
            lines.append(f"- Document digest: `{source.document_digest}`")
        if source.evidence_digest:
            lines.append(f"- Evidence digest: `{source.evidence_digest}`")
        lines.append("")
    if not analysis.sources:
        lines.extend(["Нет источников.", ""])
    lines.extend(["## Ограничения и предположения", ""])
    lines.extend(_md_bullets(analysis.risks_and_assumptions))
    lines.extend(
        [
            "",
            "---",
            "",
            "Экспорт сформирован из структурированного объекта анализа. "
            "Диалог с консультантом в файл не включён.",
            "",
        ]
    )
    return "\n".join(lines)


def render_report_markdown(report: HumanReviewedReportV1) -> str:
    lines = [
        f"# {_md_text(report.title)}",
        "",
        f"**Компания:** {_md_text(report.canonical_name)}  ",
        f"**Report ID:** `{report.id}`  ",
        f"**Версия:** {report.version}  ",
        f"**Snapshot:** {_md_text(report.as_of.isoformat())}  ",
        f"**Сформирован:** {_md_text(report.generated_at.isoformat())}  ",
        "**Статус:** human-reviewed; client_release_ready=true",
        "",
        "## Сводка выпуска",
        "",
        f"- Выпущено claims: {report.summary.released_claims}",
        f"- Элементов доказательств: {report.summary.evidence_items}",
        f"- Закрыто критических пробелов: "
        f"{report.summary.closed_critical_gaps}/6",
        f"- УДП: {report.summary.achieved_sufficiency.value} "
        f"(target {report.summary.target_sufficiency.value})",
        f"- Идентичность: {report.summary.identity_state.value}",
        f"- Целостность исполнения: "
        f"{report.summary.execution_integrity.value}",
        f"- Полнота профиля: {report.summary.profile_completeness:.0%}",
        f"- Качество evidence: {report.summary.evidence_quality:.0%}",
        f"- Коммерческий приоритет: "
        f"{report.summary.commercial_priority}/100",
        "",
    ]
    for section in report.sections:
        lines.extend([f"## {_md_text(section.title)}", ""])
        _append_markdown_table(
            lines,
            ("Показатель", "Значение", "Период", "Claim", "Evidence"),
            (
                (
                    field.predicate,
                    _display_value(field.value),
                    field.period or "—",
                    field.claim_id,
                    ", ".join(field.evidence_ids) or "—",
                )
                for field in section.fields
            ),
        )
    lines.extend(["## Критические пробелы", ""])
    _append_markdown_table(
        lines,
        ("Код", "Статус", "Закрыт", "Claims", "Reason codes"),
        (
            (
                item.code.value,
                item.status.value,
                "да" if item.closed else "нет",
                ", ".join(item.claim_ids) or "—",
                ", ".join(item.reason_codes) or "—",
            )
            for item in report.critical_gap_assessments
        ),
    )
    lines.extend(["## Приложение доказательств", ""])
    for item in report.evidence_appendix:
        lines.extend(
            [
                f"### {_md_text(item.document_title)}",
                "",
                f"**Источник:** {_md_text(item.publisher)}  ",
                f"**URL:** {_markdown_url(item.document_url)}  ",
                f"**Проверено:** {_md_text(item.document_accessed_at.isoformat())}  ",
                f"**Source kind:** {_md_text(item.source_kind.value)}",
                "",
            ]
        )
        lines.extend(
            f"> {_md_text(line)}" for line in item.quote.splitlines() or [""]
        )
        lines.extend(
            [
                "",
                f"**Locator:** {_md_text(item.locator)}  ",
                f"**Document digest:** `{item.document_digest}`  ",
                f"**Evidence digest:** `{item.evidence_digest}`",
                "",
            ]
        )
        _append_markdown_table(
            lines,
            ("Claim", "Tier", "Свежесть"),
            (
                (
                    assessment.claim_id,
                    assessment.tier.value,
                    assessment.freshness.value,
                )
                for assessment in item.claim_assessments
            ),
        )
    lines.extend(
        [
            "## Human sign-off",
            "",
            f"**Проверил:** {_md_text(report.human_sign_off.reviewer_ref)}  ",
            "**Решение:** approved  ",
            f"**Время:** "
            f"{_md_text(report.human_sign_off.decided_at.isoformat())}  ",
            f"**Основание:** {_md_text(report.human_sign_off.reason)}  ",
            f"**Sign-off digest:** "
            f"`{report.human_sign_off.sign_off_digest}`",
            "",
            "## Контроль целостности",
            "",
            f"- Канонизация: `{report.integrity.canonicalization}`",
            f"- Profile: `{report.integrity.profile_digest}`",
            f"- Evidence appendix: "
            f"`{report.integrity.evidence_appendix_digest}`",
            f"- Release control: "
            f"`{report.integrity.release_control_digest}`",
            f"- Sign-off: `{report.integrity.sign_off_digest}`",
            f"- Report: `{report.integrity.report_content_digest}`",
            "",
        ]
    )
    return "\n".join(lines)


def _set_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_document(document: Document, *, header_text: str) -> None:
    section = document.sections[0]
    _configure_section_layout(section)
    document.settings.odd_and_even_pages_header_footer = False

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)

    heading_tokens = {
        "Heading 1": (16, ACCENT, 16, 8),
        "Heading 2": (13, ACCENT, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(3)
    _set_font(header.add_run(header_text), size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(footer.add_run("AIMETON · стр. "), size=9, color=MUTED)
    _append_page_field(footer)


def _configure_section_layout(section) -> None:
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = False


def _start_new_page(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section_layout(section)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True


def _append_page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    _set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in TABLE_CELL_MARGIN_DXA.items():
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        properties.append(header)
    header.set(qn("w:val"), "true")


def _set_row_cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = properties.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        properties.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def _set_table_geometry(table: Table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != REPORT_WIDTH_DXA:
        raise ValueError("DOCX table widths must sum to 9360 DXA")
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(REPORT_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for item_width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(item_width))
        grid.append(column)

    for row in table.rows:
        _set_row_cant_split(row)
        for index, (cell, item_width) in enumerate(
            zip(row.cells, widths_dxa, strict=True)
        ):
            cell.width = Inches(item_width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(item_width))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                if index > 0 and item_width <= 1400:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_font(run, size=9.5)


def _remove_table_borders(table: Table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def _add_table(
    document: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[Any]],
    widths_dxa: Sequence[int],
) -> Table:
    prepared = list(rows)
    if not prepared:
        prepared = [
            ("Нет проверенных данных.",) + ("—",) * (len(headers) - 1)
        ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for cell, label in zip(header_row.cells, headers, strict=True):
        cell.text = str(label)
        _set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            _set_font(run, size=9.5, bold=True, color=DARK_BLUE)
    _set_repeat_table_header(header_row)
    for values in prepared:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = _display_value(value)
    _set_table_geometry(table, widths_dxa)
    document.add_paragraph()
    return table


def _add_key_value(
    document: Document,
    label: str,
    value: Any,
    *,
    code: bool = False,
) -> Paragraph:
    paragraph = document.add_paragraph()
    _set_font(paragraph.add_run(f"{label}: "), bold=True)
    run = paragraph.add_run(_display_value(value))
    _set_font(run, name="Consolas" if code else "Calibri", size=9 if code else 11)
    return paragraph


def _add_report_field(document: Document, field) -> None:
    title = document.add_paragraph()
    title.paragraph_format.keep_with_next = True
    title.paragraph_format.space_after = Pt(2)
    _set_font(title.add_run(field.predicate), bold=True, color=DARK_BLUE)

    value = document.add_paragraph(_display_value(field.value))
    value.paragraph_format.keep_with_next = True
    value.paragraph_format.space_after = Pt(2)

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(8)
    metadata_text = (
        f"Период: {field.period or '—'}  ·  Claim: {field.claim_id}  ·  "
        f"Evidence: {', '.join(field.evidence_ids) or '—'}"
    )
    _set_font(metadata.add_run(metadata_text), size=9.5, color=MUTED)


def _add_callout(document: Document, title: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT_FILL)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), "1F4D78")
        borders.append(border)
    properties.append(borders)
    title_run = paragraph.add_run(title)
    _set_font(title_run, bold=True, color=DARK_BLUE)
    title_run.add_break()
    _set_font(paragraph.add_run(text))


def _add_hyperlink(paragraph: Paragraph, url: str, label: str | None = None) -> None:
    safe_url = _safe_http_url(url)
    if safe_url is None:
        _set_font(paragraph.add_run(label or url))
        return
    relationship = paragraph.part.relate_to(
        safe_url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label or safe_url
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _document_bytes(document: Document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_site_analysis_docx(analysis: SiteAnalysis) -> bytes:
    document = Document()
    _configure_document(
        document,
        header_text="AIMETON Site Auditor · Предварительный анализ",
    )
    document.core_properties.title = (
        f"AIMETON — предварительный анализ: {analysis.company_name}"
    )
    document.core_properties.subject = "Структурированный экспорт анализа сайта"
    document.core_properties.author = "AIMETON"

    document.add_paragraph(
        "AIMETON — ПРЕДВАРИТЕЛЬНЫЙ АНАЛИЗ КОМПАНИИ",
        style="Title",
    )
    document.add_paragraph(analysis.company_name, style="Subtitle")
    _add_callout(
        document,
        "Статус",
        "Предварительный результат старого контура анализа. "
        "Он не является подписанным Report v1 и не прошёл human sign-off.",
    )
    url_paragraph = document.add_paragraph()
    _set_font(url_paragraph.add_run("Сайт: "), bold=True)
    _add_hyperlink(url_paragraph, analysis.url)
    document.add_paragraph(analysis.business_summary)

    document.add_heading("Состояние допустимости", level=1)
    _add_table(
        document,
        ("Показатель", "Значение"),
        (
            ("Analysis state", analysis.readiness.analysis_state),
            (
                "Client release",
                str(analysis.readiness.client_release_eligible).lower(),
            ),
            ("УДП", analysis.readiness.sufficiency_level),
            ("Identity", analysis.readiness.identity_state),
            (
                "Полнота профиля",
                f"{analysis.readiness.profile_completeness:.0%}",
            ),
            (
                "Качество evidence",
                f"{analysis.readiness.evidence_quality:.0%}",
            ),
            (
                "Коммерческий приоритет",
                f"{analysis.readiness.commercial_priority}/100",
            ),
            ("Budget", analysis.readiness.budget_state),
            (
                "Providers",
                ", ".join(
                    f"{key}={value}"
                    for key, value in sorted(
                        analysis.readiness.provider_states.items()
                    )
                )
                or "not_reported",
            ),
            (
                "Обязательные вертикали",
                ", ".join(
                    f"{item.code}={item.state}"
                    for item in analysis.readiness.required_verticals
                )
                or "not_reported",
            ),
            (
                "Блокеры",
                ", ".join(analysis.readiness.release_blockers) or "—",
            ),
        ),
        (2700, 6660),
    )

    opportunity = analysis.commercial_opportunity
    document.add_heading("Коммерческая возможность", level=1)
    _add_table(
        document,
        ("Показатель", "Значение"),
        (
            ("Тип", opportunity.opportunity_type),
            ("Оценка", f"{opportunity.score}/100"),
            ("Квалификация", opportunity.qualification),
            ("Гипотеза проблемы", opportunity.problem_hypothesis),
            ("Рекомендуемое решение", opportunity.recommended_solution),
            ("Ожидаемая ценность", opportunity.expected_value),
        ),
        (2700, 6660),
    )

    document.add_heading("Экономические сигналы", level=1)
    _add_table(
        document,
        ("Сигнал", "Основание", "Эффект", "Уверенность", "Источники"),
        (
            (
                item.signal,
                item.evidence,
                item.business_effect,
                item.confidence,
                ", ".join(item.source_ids) or "—",
            )
            for item in analysis.economic_signals
        ),
        (1800, 2300, 2300, 1300, 1660),
    )

    document.add_heading("Факты о компании", level=1)
    _add_table(
        document,
        ("Поле", "Значение", "Период", "Уверенность", "Источники"),
        (
            (
                item.field,
                item.value,
                item.period or "—",
                item.confidence,
                ", ".join(item.source_ids) or "—",
            )
            for item in analysis.company_facts
        ),
        (1650, 3350, 1200, 1350, 1810),
    )

    document.add_heading("Бизнес-машина AIMETON 4×4", level=1)
    for item in analysis.business_machine_4x4:
        document.add_heading(f"{item.code} · {item.vertex}", level=2)
        _add_key_value(document, "Оператор", item.detail_operator)
        _add_key_value(
            document,
            "Статус",
            f"{item.status}; уверенность: {item.confidence}",
        )
        _add_key_value(
            document,
            "Источники",
            ", ".join(item.source_ids) or "—",
        )
        document.add_paragraph(item.finding)
        if item.sales_relevance:
            _add_callout(
                document,
                "Значение для продажи",
                item.sales_relevance,
            )
    if not analysis.business_machine_4x4:
        document.add_paragraph("Нет данных.")

    document.add_heading("Подходящие AI-инструменты", level=1)
    _add_table(
        document,
        ("Инструмент", "Назначение", "Польза", "Приоритет"),
        (
            (item.name, item.purpose, item.benefit, item.priority)
            for item in analysis.agents
        ),
        (1900, 2800, 3300, 1360),
    )

    action = analysis.action_package
    document.add_heading("Пакет действия", level=1)
    _add_key_value(
        document,
        "Предполагаемый ЛПР",
        action.decision_maker_hypothesis,
    )
    _add_key_value(document, "Основание для контакта", action.contact_reason)
    document.add_heading("Демонстрационный сценарий", level=2)
    if action.demo_scenario:
        for step in action.demo_scenario:
            document.add_paragraph(step, style="List Number")
    else:
        document.add_paragraph("Нет данных.")
    _add_callout(document, "Первое сообщение", action.first_message)
    _add_key_value(document, "Следующий шаг", action.next_action)

    document.add_heading("Источники", level=1)
    for source in analysis.sources:
        document.add_heading(f"{source.id} · {source.title}", level=2)
        paragraph = document.add_paragraph()
        _set_font(paragraph.add_run("URL: "), bold=True)
        _add_hyperlink(paragraph, source.url)
        _add_key_value(document, "Проверено", source.accessed_at)
        _add_key_value(document, "Тип", source.source_type)
        _add_key_value(document, "Уровень", source.evidence_level)
        _add_callout(document, "Цитата", source.evidence_quote)
        if source.document_digest:
            _add_key_value(
                document,
                "Document digest",
                source.document_digest,
                code=True,
            )
        if source.evidence_digest:
            _add_key_value(
                document,
                "Evidence digest",
                source.evidence_digest,
                code=True,
            )
    if not analysis.sources:
        document.add_paragraph("Нет источников.")

    document.add_heading("Ограничения и предположения", level=1)
    if analysis.risks_and_assumptions:
        for item in analysis.risks_and_assumptions:
            document.add_paragraph(item, style="List Bullet")
    else:
        document.add_paragraph("Нет данных.")
    _add_callout(
        document,
        "Состав экспорта",
        "Файл сформирован из структурированного объекта анализа. "
        "Диалог с консультантом в него не включён.",
    )
    return _document_bytes(document)


def render_report_docx(report: HumanReviewedReportV1) -> bytes:
    document = Document()
    _configure_document(
        document,
        header_text="AIMETON Report v1 · Human reviewed",
    )
    document.core_properties.title = report.title
    document.core_properties.subject = "Human-reviewed evidence report"
    document.core_properties.author = "AIMETON"
    document.core_properties.created = report.generated_at.replace(tzinfo=None)
    document.core_properties.modified = report.generated_at.replace(tzinfo=None)

    document.add_paragraph(report.title.upper(), style="Title")
    document.add_paragraph(report.canonical_name, style="Subtitle")
    _add_callout(
        document,
        "Статус выпуска",
        "Human-reviewed; client_release_ready=true. "
        "Содержимое связано с точным snapshot профиля и приложения доказательств.",
    )
    _add_table(
        document,
        ("Параметр", "Значение"),
        (
            ("Report ID", report.id),
            ("Версия", report.version),
            ("Snapshot", report.as_of.isoformat()),
            ("Сформирован", report.generated_at.isoformat()),
            ("Mission ID", report.mission_id),
            ("Entity ID", report.entity_id),
            ("Profile ID", report.profile_id),
        ),
        (2700, 6660),
    )

    document.add_heading("Сводка выпуска", level=1)
    _add_table(
        document,
        ("Показатель", "Значение"),
        (
            ("Выпущено claims", report.summary.released_claims),
            ("Элементов доказательств", report.summary.evidence_items),
            (
                "Закрыто критических пробелов",
                f"{report.summary.closed_critical_gaps}/6",
            ),
            (
                "УДП",
                f"{report.summary.achieved_sufficiency.value} "
                f"(target {report.summary.target_sufficiency.value})",
            ),
            ("Идентичность", report.summary.identity_state.value),
            (
                "Целостность исполнения",
                report.summary.execution_integrity.value,
            ),
            (
                "Полнота профиля",
                f"{report.summary.profile_completeness:.0%}",
            ),
            (
                "Качество evidence",
                f"{report.summary.evidence_quality:.0%}",
            ),
            (
                "Коммерческий приоритет",
                f"{report.summary.commercial_priority}/100",
            ),
        ),
        (4700, 4660),
    )

    _start_new_page(document)
    for section in report.sections:
        document.add_heading(section.title, level=1)
        if section.fields:
            for field in section.fields:
                _add_report_field(document, field)
        else:
            paragraph = document.add_paragraph(
                "Проверенных данных для выпуска нет."
            )
            for run in paragraph.runs:
                _set_font(run, italic=True, color=MUTED)

    document.add_heading("Критические пробелы", level=1)
    _add_table(
        document,
        ("Код", "Статус", "Закрыт", "Claims", "Reason codes"),
        (
            (
                item.code.value,
                item.status.value,
                "да" if item.closed else "нет",
                ", ".join(item.claim_ids) or "—",
                ", ".join(item.reason_codes) or "—",
            )
            for item in report.critical_gap_assessments
        ),
        (1450, 1350, 1000, 2780, 2780),
    )

    document.add_heading("Приложение доказательств", level=1)
    container_table = document.add_table(rows=0, cols=1)
    for item in report.evidence_appendix:
        container = container_table.add_row().cells[0]
        container.add_paragraph(item.document_title, style="Heading 2")
        _add_key_value(container, "Источник", item.publisher)
        paragraph = container.add_paragraph()
        _set_font(paragraph.add_run("URL: "), bold=True)
        _add_hyperlink(paragraph, str(item.document_url))
        _add_key_value(
            container,
            "Проверено",
            item.document_accessed_at.isoformat(),
        )
        _add_key_value(container, "Source kind", item.source_kind.value)
        _add_callout(container, "Цитата", item.quote)
        _add_key_value(container, "Locator", item.locator)
        _add_key_value(
            container,
            "Document digest",
            item.document_digest,
            code=True,
        )
        _add_key_value(
            container,
            "Evidence digest",
            item.evidence_digest,
            code=True,
        )
        _add_table(
            container,
            ("Claim", "Tier", "Свежесть"),
            (
                (
                    assessment.claim_id,
                    assessment.tier.value,
                    assessment.freshness.value,
                )
                for assessment in item.claim_assessments
            ),
            (3900, 2700, 2760),
        )
    _set_table_geometry(container_table, (REPORT_WIDTH_DXA,))
    _remove_table_borders(container_table)

    document.add_heading("Human sign-off", level=1)
    _add_key_value(
        document,
        "Проверил",
        report.human_sign_off.reviewer_ref,
    )
    _add_key_value(document, "Решение", "approved")
    _add_key_value(
        document,
        "Время",
        report.human_sign_off.decided_at.isoformat(),
    )
    _add_key_value(document, "Основание", report.human_sign_off.reason)
    _add_key_value(
        document,
        "Sign-off digest",
        report.human_sign_off.sign_off_digest,
        code=True,
    )

    document.add_heading("Контроль целостности", level=1)
    _add_key_value(
        document,
        "Канонизация",
        report.integrity.canonicalization,
        code=True,
    )
    _add_key_value(
        document,
        "Profile",
        report.integrity.profile_digest,
        code=True,
    )
    _add_key_value(
        document,
        "Evidence appendix",
        report.integrity.evidence_appendix_digest,
        code=True,
    )
    _add_key_value(
        document,
        "Release control",
        report.integrity.release_control_digest,
        code=True,
    )
    _add_key_value(
        document,
        "Sign-off",
        report.integrity.sign_off_digest,
        code=True,
    )
    _add_key_value(
        document,
        "Report",
        report.integrity.report_content_digest,
        code=True,
    )
    return _document_bytes(document)
