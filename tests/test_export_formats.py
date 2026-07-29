from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.heuristics import heuristic_analysis
from app.main import app
from app.models import CompanyFact, EvidenceSource
from app.scraper import extract_visible_text
from app.sef.exports import (
    render_site_analysis_docx,
    render_site_analysis_markdown,
)


ROOT = Path(__file__).resolve().parents[1]


def _analysis():
    html = (ROOT / "tests" / "fixtures" / "dns.html").read_text(
        encoding="utf-8"
    )
    title, text = extract_visible_text(html)
    result = heuristic_analysis("https://dns.example", title, text)
    return result.model_copy(
        update={
            "company_name": "DNS <script>alert('x')</script>",
            "company_facts": [
                CompanyFact(
                    field="website",
                    value="https://dns.example",
                    confidence="Высокая",
                    source_ids=["S1"],
                )
            ],
            "sources": [
                EvidenceSource(
                    id="S1",
                    title="Официальный сайт",
                    url="https://dns.example",
                    accessed_at="2026-07-29T06:00:00+02:00",
                    evidence_quote="Каталог товаров и сборка ПК.",
                    source_type="official_page",
                    evidence_level="confirmed_fact",
                )
            ],
            "risks_and_assumptions": [
                "Предварительный результат; требуется human review."
            ],
        }
    )


def _docx_text(payload: bytes) -> str:
    document = Document(io.BytesIO(payload))
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_preliminary_markdown_is_structured_and_has_no_raw_html_or_chat():
    rendered = render_site_analysis_markdown(_analysis())

    assert "# AIMETON — предварительный анализ компании" in rendered
    assert "## Факты о компании" in rendered
    assert "## Источники" in rendered
    assert "human sign-off" in rendered
    assert "<script>" not in rendered
    assert "&lt;script&gt;" in rendered
    assert "Диалог с консультантом в файл не включён." in rendered


def test_preliminary_docx_is_native_editable_word_content():
    payload = render_site_analysis_docx(_analysis())
    text = _docx_text(payload)

    assert payload.startswith(b"PK")
    assert "ПРЕДВАРИТЕЛЬНЫЙ АНАЛИЗ КОМПАНИИ" in text
    assert "Факты о компании" in text
    assert "Официальный сайт" in text
    assert "Диалог с консультантом" in text
    assert "<script>" in text


def test_preliminary_export_endpoints_return_attachments():
    client = TestClient(app)
    request = _analysis().model_dump(mode="json")

    markdown = client.post("/api/export/analysis.md", json=request)
    word = client.post("/api/export/analysis.docx", json=request)

    assert markdown.status_code == 200
    assert markdown.headers["content-type"].startswith("text/markdown")
    assert markdown.headers["content-disposition"].endswith(
        'filename="aimeton-preliminary-analysis.md"'
    )
    assert word.status_code == 200
    assert word.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert word.headers["content-disposition"].endswith(
        'filename="aimeton-preliminary-analysis.docx"'
    )
    assert word.content.startswith(b"PK")


def test_ui_exports_and_chat_are_scoped_to_active_analysis():
    script = (ROOT / "static" / "app.js").read_text(encoding="utf-8")

    assert "const CHAT_KEY = 'aimeton_chat_sessions'" in script
    assert "ui_analysis_id" in script
    assert "renderChatSession()" in script
    assert "currentChatSession().slice(-12)" in script
    assert "replacedIds.forEach(id => delete sessions[id])" in script
    assert "/api/export/analysis.md" in script
    assert "/api/export/analysis.docx" in script
    assert "msgNodes" not in script
