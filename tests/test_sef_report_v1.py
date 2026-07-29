from __future__ import annotations

import copy
import io
import json
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from jsonschema import Draft202012Validator

from app.main import app
from app.sef.ledger import LedgerRequest
from app.sef.models import SefBundle
from app.sef.release_control import MissionReleaseControl
from app.sef.exports import render_report_docx, render_report_markdown
from app.sef.report import (
    ReportBuildRequest,
    ReportReleaseError,
    ReportReviewPackageRequest,
    build_human_reviewed_report,
    build_review_package_from_request,
    render_report_html,
)
from scripts.export_sef_report_schema import TARGET, render_schema


ROOT = Path(__file__).resolve().parents[1]
FIXTURE = ROOT / "tests" / "fixtures" / "sef" / "positive-chain-v0.1.json"


def _digest(seed: int) -> str:
    return f"sha256:{seed:064x}"


def _metadata(evidence_id: str, seed: int) -> dict:
    return {
        "id": f"evidence_meta_{seed}",
        "mission_id": "mission_golden_01",
        "evidence_id": evidence_id,
        "correlation_id": "corr_golden_01",
        "tier": "tier_2_first_party",
        "valid_at": "2026-07-01T00:00:00Z",
        "fresh_until": "2027-07-01T00:00:00Z",
        "recorded_at": "2026-07-28T10:05:00Z",
    }


def _complete_payload() -> tuple[dict, list[dict]]:
    payload = json.loads(FIXTURE.read_text(encoding="utf-8"))
    metadata = [_metadata("evidence_golden_01", 1)]
    specs = [
        ("email", "info@example.ru", "contact", 10),
        ("headcount", 125, "workforce", 20),
        (
            "revenue",
            {"amount": "120000000", "currency": "RUB", "period": "2025"},
            "finance",
            30,
        ),
        ("beneficial_owner", "Иванов Иван Иванович", "ownership", 40),
        ("affiliate", "ООО «Связанная компания»", "affiliation", 50),
        (
            "legal_events_summary",
            "Проверены арбитраж, взыскания и банкротство за 2025–2026 годы",
            "legal",
            60,
        ),
    ]
    for predicate, value, source_label, seed in specs:
        source_id = f"source_{seed}"
        document_id = f"document_{seed}"
        evidence_id = f"evidence_{seed}"
        claim_id = f"claim_{seed}"
        payload["sources"].append(
            {
                "id": source_id,
                "mission_id": "mission_golden_01",
                "correlation_id": "corr_golden_01",
                "kind": (
                    "official_registry"
                    if source_label in {"finance", "legal"}
                    else "first_party"
                ),
                "publisher": f"Проверяемый источник {source_label}",
                "homepage_url": f"https://example.ru/{source_label}",
                "terms_ref": None,
            }
        )
        payload["documents"].append(
            {
                "id": document_id,
                "mission_id": "mission_golden_01",
                "source_id": source_id,
                "correlation_id": "corr_golden_01",
                "url": f"https://example.ru/{source_label}/document",
                "title": f"Документ {source_label}",
                "accessed_at": "2026-07-28T10:02:00Z",
                "fetch_status": "fetched",
                "content_digest": _digest(seed),
                "media_type": "text/html",
            }
        )
        payload["evidence"].append(
            {
                "id": evidence_id,
                "mission_id": "mission_golden_01",
                "source_id": source_id,
                "document_id": document_id,
                "correlation_id": "corr_golden_01",
                "evidence_type": "document_quote",
                "quote": f"Проверяемое значение {predicate}: {value}",
                "locator": f"section/{predicate}",
                "observed_at": "2026-07-28T10:03:00Z",
                "digest": _digest(seed + 1),
            }
        )
        payload["claims"].append(
            {
                "id": claim_id,
                "mission_id": "mission_golden_01",
                "entity_id": "entity_golden_01",
                "correlation_id": "corr_golden_01",
                "predicate": predicate,
                "value": value,
                "state": "confirmed",
                "critical": True,
                "evidence_refs": [
                    {"evidence_id": evidence_id, "relation": "supports"}
                ],
                "created_at": "2026-07-28T10:03:30Z",
            }
        )
        payload["review_decisions"].append(
            {
                "id": f"review_{seed}",
                "mission_id": "mission_golden_01",
                "correlation_id": "corr_golden_01",
                "target_type": "claim",
                "target_id": claim_id,
                "decision": "approved",
                "reviewer_ref": "reviewer_human_01",
                "reason": "Документ и значение проверены оператором.",
                "decided_at": "2026-07-28T10:04:00Z",
            }
        )
        metadata.append(_metadata(evidence_id, seed))
    return payload, metadata


def _release_control(**overrides) -> MissionReleaseControl:
    payload = {
        "mission_id": "mission_golden_01",
        "evaluated_at": "2026-07-28T12:10:00Z",
        "target_sufficiency": "L4",
        "achieved_sufficiency": "L4",
        "identity_state": "resolved",
        "execution_integrity": "validated",
        "analysis_state": "schema_validated",
        "unresolved_critical_conflicts": 0,
        "required_verticals": [
            {"code": code, "required": True, "state": "verified"}
            for code in (
                "identity",
                "contacts",
                "workforce",
                "financials",
                "ownership",
                "legal_events",
            )
        ],
        "providers": [
            {
                "provider_ref": "yandex",
                "required": True,
                "state": "active",
            },
            {
                "provider_ref": "searxng",
                "required": False,
                "state": "active",
            },
        ],
        "budget_state": "within_budget",
        "profile_completeness": 0.82,
        "evidence_quality": 0.91,
        "commercial_priority": 67,
        "reason_codes": [],
    }
    payload.update(overrides)
    return MissionReleaseControl.model_validate(payload)


def _source_request(
    payload: dict | None = None,
    metadata: list[dict] | None = None,
    release_control: MissionReleaseControl | None = None,
) -> ReportReviewPackageRequest:
    if payload is None or metadata is None:
        payload, metadata = _complete_payload()
    return ReportReviewPackageRequest(
        bundle=SefBundle.model_validate(payload),
        ledger_request=LedgerRequest.model_validate(
            {
                "mission_id": "mission_golden_01",
                "as_of": "2026-07-28T12:00:00Z",
                "evidence_metadata": metadata,
            }
        ),
        entity_id="entity_golden_01",
        release_control=release_control or _release_control(),
    )


def _report_request(
    source: ReportReviewPackageRequest | None = None,
    **review_overrides,
) -> ReportBuildRequest:
    source = source or _source_request()
    package = build_review_package_from_request(source)
    review = {
        "reviewer_ref": "reviewer_human_01",
        "decision": "approved",
        "reason": "Профиль и приложение доказательств проверены.",
        "decided_at": "2026-07-28T12:30:00Z",
        "attested": True,
        "reviewed_profile_digest": package.profile_digest,
        "reviewed_evidence_appendix_digest": package.evidence_appendix_digest,
        "reviewed_release_control_digest": package.release_control_digest,
    }
    review.update(review_overrides)
    return ReportBuildRequest(
        **source.model_dump(),
        review=review,
        generated_at="2026-07-28T13:00:00Z",
        title="Доказательный профиль компании",
        version=1,
    )


def test_report_schema_is_current_and_valid():
    assert TARGET.read_text(encoding="utf-8") == render_schema()
    schema = json.loads(TARGET.read_text(encoding="utf-8"))
    Draft202012Validator.check_schema(schema)
    assert schema["$id"].endswith("/sef-report-v1.schema.json")


def test_review_package_binds_exact_profile_and_evidence_snapshot():
    source = _source_request()
    package = build_review_package_from_request(source)

    assert package.reviewable is True
    assert package.blockers == []
    assert package.profile.summary.client_release_ready is False
    assert package.profile_digest.startswith("sha256:")
    assert package.evidence_appendix_digest.startswith("sha256:")
    assert package.release_control_digest.startswith("sha256:")
    assert package.release_control.achieved_sufficiency.value == "L4"

    response = TestClient(app).post(
        "/api/sef/report/review-package",
        json=source.model_dump(mode="json"),
    )
    assert response.status_code == 200
    assert response.json()["profile_digest"] == package.profile_digest


def test_approved_review_releases_deterministic_report():
    request = _report_request()
    first = build_human_reviewed_report(request)
    second = build_human_reviewed_report(request)

    assert first == second
    assert first.id == second.id
    assert first.summary.client_release_ready is True
    assert first.summary.profile_gate_passed is True
    assert first.summary.human_review_approved is True
    assert first.summary.closed_critical_gaps == 6
    assert first.summary.achieved_sufficiency.value == "L4"
    assert first.summary.identity_state.value == "resolved"
    assert first.summary.profile_completeness == 0.82
    assert first.summary.evidence_quality == 0.91
    assert first.summary.commercial_priority == 67
    assert first.claim_ids == sorted(first.claim_ids)
    assert all(
        field.client_eligible and field.status.value == "verified"
        for section in first.sections
        for field in section.fields
    )
    assert {
        assessment.claim_id
        for item in first.evidence_appendix
        for assessment in item.claim_assessments
    } <= set(first.claim_ids)


def test_report_build_does_not_mutate_inputs():
    request = _report_request()
    before = request.model_dump(mode="json")

    build_human_reviewed_report(request)

    assert request.model_dump(mode="json") == before


@pytest.mark.parametrize(
    ("field", "value", "blocker"),
    [
        (
            "reviewed_profile_digest",
            _digest(900),
            "review_profile_digest_mismatch",
        ),
        (
            "reviewed_evidence_appendix_digest",
            _digest(901),
            "review_evidence_appendix_digest_mismatch",
        ),
        (
            "reviewed_release_control_digest",
            _digest(902),
            "review_release_control_digest_mismatch",
        ),
        ("decision", "rejected", "human_review_not_approved"),
        ("attested", False, "human_attestation_missing"),
        (
            "decided_at",
            "2026-07-28T11:59:59Z",
            "human_review_predates_profile_snapshot",
        ),
        (
            "decided_at",
            "2026-07-28T13:00:01Z",
            "human_review_after_report_generation",
        ),
    ],
)
def test_report_release_fails_closed(field: str, value, blocker: str):
    request = _report_request(**{field: value})

    with pytest.raises(ReportReleaseError) as exc:
        build_human_reviewed_report(request)

    assert blocker in exc.value.blockers


def test_profile_with_hypothesis_is_not_reviewable_or_releasable():
    payload, metadata = _complete_payload()
    claim = next(item for item in payload["claims"] if item["predicate"] == "affiliate")
    claim["state"] = "candidate"
    payload["review_decisions"] = [
        item
        for item in payload["review_decisions"]
        if item["target_id"] != claim["id"]
    ]
    source = _source_request(payload, metadata)
    package = build_review_package_from_request(source)

    assert package.reviewable is False
    assert "company_profile_gate_not_passed" in package.blockers
    assert "profile_contains_hypothesis" in package.blockers
    with pytest.raises(ReportReleaseError):
        build_human_reviewed_report(_report_request(source))


@pytest.mark.parametrize(
    ("overrides", "blocker"),
    [
        (
            {"execution_integrity": "degraded"},
            "execution_integrity_degraded",
        ),
        (
            {"execution_integrity": "validation_error"},
            "execution_integrity_validation_error",
        ),
        (
            {"analysis_state": "preliminary_hypothesis"},
            "analysis_preliminary_hypothesis",
        ),
        ({"analysis_state": "validation_error"}, "analysis_validation_error"),
        ({"identity_state": "unresolved"}, "identity_unresolved"),
        ({"identity_state": "conflicting"}, "identity_conflicting"),
        (
            {"mission_id": "mission_other"},
            "release_control_mission_mismatch",
        ),
        ({"target_sufficiency": "L3"}, "target_sufficiency_below_l4"),
        ({"achieved_sufficiency": "L3"}, "sufficiency_below_l4"),
        (
            {"unresolved_critical_conflicts": 1},
            "unresolved_critical_conflict",
        ),
        (
            {
                "required_verticals": [
                    {
                        "code": "identity",
                        "required": True,
                        "state": "degraded",
                    }
                ]
            },
            "required_vertical_identity_degraded",
        ),
        (
            {
                "providers": [
                    {
                        "provider_ref": "yandex",
                        "required": True,
                        "state": "pricing_unknown",
                    }
                ]
            },
            "required_provider_yandex_pricing_unknown",
        ),
        ({"budget_state": "unknown"}, "budget_unknown"),
        ({"budget_state": "exhausted"}, "budget_exhausted"),
    ],
)
def test_mission_release_control_blocks_unsafe_client_release(
    overrides: dict,
    blocker: str,
):
    source = _source_request(release_control=_release_control(**overrides))
    package = build_review_package_from_request(source)

    assert package.reviewable is False
    assert blocker in package.blockers
    with pytest.raises(ReportReleaseError) as exc:
        build_human_reviewed_report(_report_request(source))
    assert blocker in exc.value.blockers


def test_missing_release_control_is_rejected_before_review():
    source = _source_request().model_dump(mode="json")
    source.pop("release_control")

    response = TestClient(app).post(
        "/api/sef/report/review-package",
        json=source,
    )

    assert response.status_code == 422


def test_report_api_recalculates_snapshot_and_returns_blockers():
    request = _report_request(reviewed_profile_digest=_digest(999))
    response = TestClient(app).post(
        "/api/sef/report",
        json=request.model_dump(mode="json"),
    )

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "report_release_blocked"
    assert (
        "review_profile_digest_mismatch"
        in response.json()["detail"]["blockers"]
    )

    forged = request.model_dump(mode="json")
    forged["client_release_ready"] = True
    assert TestClient(app).post("/api/sef/report", json=forged).status_code == 422


def test_html_export_is_escaped_and_contains_no_search_snippet():
    payload, metadata = _complete_payload()
    payload["discovery_hints"][0]["snippet"] = "FORBIDDEN_SEARCH_SNIPPET"
    contact = next(item for item in payload["claims"] if item["predicate"] == "email")
    contact["value"] = "<script>alert('x')</script>"
    report = build_human_reviewed_report(
        _report_request(_source_request(payload, metadata))
    )

    rendered = render_report_html(report)

    assert "FORBIDDEN_SEARCH_SNIPPET" not in rendered
    assert "<script>alert('x')</script>" not in rendered
    assert "&lt;script&gt;" in rendered
    assert report.integrity.report_content_digest in rendered

    response = TestClient(app).post(
        "/api/sef/report.html",
        json=_report_request(_source_request(payload, metadata)).model_dump(
            mode="json"
        ),
    )
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/html")
    assert "FORBIDDEN_SEARCH_SNIPPET" not in response.text


def test_markdown_and_docx_exports_are_bound_to_signed_report():
    report = build_human_reviewed_report(_report_request())
    markdown = render_report_markdown(report)
    word = render_report_docx(report)
    word_document = Document(io.BytesIO(word))
    word_text = "\n".join(
        [
            *(paragraph.text for paragraph in word_document.paragraphs),
            *(
                cell.text
                for table in word_document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )

    assert report.id in markdown
    assert report.integrity.report_content_digest in markdown
    assert "Критические пробелы" in markdown
    assert word.startswith(b"PK")
    assert report.id in word_text
    assert report.integrity.report_content_digest in word_text
    assert "Human sign-off" in word_text


@pytest.mark.parametrize(
    ("path", "extension", "media_type"),
    [
        ("/api/sef/report.md", "md", "text/markdown"),
        (
            "/api/sef/report.docx",
            "docx",
            "application/vnd.openxmlformats-officedocument",
        ),
    ],
)
def test_signed_export_endpoints_are_attachments(
    path: str,
    extension: str,
    media_type: str,
):
    request = _report_request()
    report = build_human_reviewed_report(request)
    response = TestClient(app).post(path, json=request.model_dump(mode="json"))

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(media_type)
    assert response.headers["content-disposition"].endswith(
        f'filename="aimeton-{report.id}.{extension}"'
    )
    assert response.headers["x-aimeton-report-id"] == report.id
    assert (
        response.headers["x-aimeton-report-digest"]
        == report.integrity.report_content_digest
    )


@pytest.mark.parametrize(
    "path",
    ["/api/sef/report.md", "/api/sef/report.docx"],
)
def test_signed_export_endpoints_fail_closed(path: str):
    request = _report_request(reviewed_profile_digest=_digest(777))
    response = TestClient(app).post(path, json=request.model_dump(mode="json"))

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "report_release_blocked"


def test_changed_profile_invalidates_previous_human_review():
    source = _source_request()
    signed = _report_request(source)
    changed_payload = copy.deepcopy(source.bundle.model_dump(mode="json"))
    changed_claim = next(
        item for item in changed_payload["claims"] if item["predicate"] == "email"
    )
    changed_claim["value"] = "new@example.ru"
    changed_source = _source_request(
        changed_payload,
        source.ledger_request.model_dump(mode="json")["evidence_metadata"],
    )
    changed_request = ReportBuildRequest(
        **changed_source.model_dump(),
        review=signed.review,
        generated_at=signed.generated_at,
        title=signed.title,
        version=signed.version,
    )

    with pytest.raises(ReportReleaseError) as exc:
        build_human_reviewed_report(changed_request)

    assert "review_profile_digest_mismatch" in exc.value.blockers


def test_changed_release_control_invalidates_previous_human_review():
    source = _source_request()
    signed = _report_request(source)
    changed_source = _source_request(
        source.bundle.model_dump(mode="json"),
        source.ledger_request.model_dump(mode="json")["evidence_metadata"],
        _release_control(commercial_priority=68),
    )
    changed_request = ReportBuildRequest(
        **changed_source.model_dump(),
        review=signed.review,
        generated_at=signed.generated_at,
        title=signed.title,
        version=signed.version,
    )

    with pytest.raises(ReportReleaseError) as exc:
        build_human_reviewed_report(changed_request)

    assert "review_release_control_digest_mismatch" in exc.value.blockers
